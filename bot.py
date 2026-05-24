import asyncio
import logging
import os
import json
import aiosqlite
from datetime import datetime
from aiogram import Bot, Dispatcher, F
from aiogram.filters import CommandStart, Command
from aiogram.types import Message, CallbackQuery, FSInputFile
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
import httpx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(level=logging.INFO)

BOT_TOKEN = os.getenv("BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
ADMIN_IDS = list(map(int, os.getenv("ADMIN_IDS", "0").split(",")))
DB = "bot.db"
DOCS_DIR = "docs"
os.makedirs(DOCS_DIR, exist_ok=True)

# ─── TILLAR ───────────────────────────────────────────────────────────────────

TEXTS = {
    "uz": {
        "welcome": "⚖️ <b>Yuridik Yordamchi Botiga Xush Kelibsiz!</b>\n\n📄 Hujjatlar konstruktori\n🤖 AI huquqiy maslahat\n\nXizmatni tanlang 👇",
        "choose_doc": "📂 <b>Hujjat turini tanlang:</b>",
        "info": "ℹ️ <b>Xizmatlar narxi:</b>\n\n• Barcha hujjatlar — 10,000 so'm\n• AI Maslahat — 10,000 so'm",
        "ai_start": "🤖 <b>AI Huquqiy Maslahat</b>\n\nHuquqiy muammoingizni yozing:",
        "thinking": "🤔 Tahlil qilinmoqda...",
        "ai_error": "❌ AI xizmati hozircha mavjud emas.",
        "paid_ok": "✅ <b>To'lov tasdiqlandi! Hujjatingiz tayyor.</b>\n\nWord (.docx) formatida yuborildi.",
        "preparing": "⏳ Hujjat tayyorlanmoqda...",
        "cancel": "❌ Bekor qilish",
        "back": "⬅️ Orqaga",
        "menu_docs": "📄 Hujjat olish",
        "menu_ai": "🤖 AI Maslahat",
        "menu_info": "ℹ️ Ma'lumot",
        "menu_lang": "🌐 Til",
        "pay_btn": "💳 To'lash (Demo)",
        "confirm_btn": "✅ Tasdiqlash (Demo)",
        "admin_title": "👨‍💼 <b>Admin Panel</b>",
        "stats_title": "📊 <b>Statistika</b>",
        "broadcast_ask": "📢 Barcha foydalanuvchilarga yuboriladigan xabarni yozing:",
        "broadcast_done": "✅ {n} ta foydalanuvchiga yuborildi.",
        "no_access": "❌ Ruxsat yo'q.",
        "choose_lang": "🌐 Tilni tanlang:",
        "lang_saved": "✅ Til saqlandi!",
    },
    "ru": {
        "welcome": "⚖️ <b>Добро пожаловать в Юридический Помощник!</b>\n\n📄 Конструктор документов\n🤖 AI юридическая консультация\n\nВыберите услугу 👇",
        "choose_doc": "📂 <b>Выберите тип документа:</b>",
        "info": "ℹ️ <b>Стоимость услуг:</b>\n\n• Все документы — 10,000 сум\n• AI консультация — 10,000 сум",
        "ai_start": "🤖 <b>AI Юридическая Консультация</b>\n\nОпишите вашу юридическую проблему:",
        "thinking": "🤔 Анализируется...",
        "ai_error": "❌ AI сервис временно недоступен.",
        "paid_ok": "✅ <b>Оплата подтверждена! Ваш документ готов.</b>\n\nОтправлен в формате Word (.docx).",
        "preparing": "⏳ Документ готовится...",
        "cancel": "❌ Отмена",
        "back": "⬅️ Назад",
        "menu_docs": "📄 Получить документ",
        "menu_ai": "🤖 AI Консультация",
        "menu_info": "ℹ️ Информация",
        "menu_lang": "🌐 Язык",
        "pay_btn": "💳 Оплатить (Демо)",
        "confirm_btn": "✅ Подтвердить (Демо)",
        "admin_title": "👨‍💼 <b>Панель администратора</b>",
        "stats_title": "📊 <b>Статистика</b>",
        "broadcast_ask": "📢 Напишите сообщение для всех пользователей:",
        "broadcast_done": "✅ Отправлено {n} пользователям.",
        "no_access": "❌ Нет доступа.",
        "choose_lang": "🌐 Выберите язык:",
        "lang_saved": "✅ Язык сохранён!",
    },
    "en": {
        "welcome": "⚖️ <b>Welcome to Legal Assistant Bot!</b>\n\n📄 Document constructor\n🤖 AI legal consultation\n\nChoose a service 👇",
        "choose_doc": "📂 <b>Select document type:</b>",
        "info": "ℹ️ <b>Service prices:</b>\n\n• All documents — 10,000 UZS\n• AI consultation — 10,000 UZS",
        "ai_start": "🤖 <b>AI Legal Consultation</b>\n\nDescribe your legal issue:",
        "thinking": "🤔 Analyzing...",
        "ai_error": "❌ AI service is temporarily unavailable.",
        "paid_ok": "✅ <b>Payment confirmed! Your document is ready.</b>\n\nSent in Word (.docx) format.",
        "preparing": "⏳ Preparing document...",
        "cancel": "❌ Cancel",
        "back": "⬅️ Back",
        "menu_docs": "📄 Get document",
        "menu_ai": "🤖 AI Consultation",
        "menu_info": "ℹ️ Information",
        "menu_lang": "🌐 Language",
        "pay_btn": "💳 Pay (Demo)",
        "confirm_btn": "✅ Confirm (Demo)",
        "admin_title": "👨‍💼 <b>Admin Panel</b>",
        "stats_title": "📊 <b>Statistics</b>",
        "broadcast_ask": "📢 Write a message for all users:",
        "broadcast_done": "✅ Sent to {n} users.",
        "no_access": "❌ Access denied.",
        "choose_lang": "🌐 Choose language:",
        "lang_saved": "✅ Language saved!",
    }
}

DOC_NAMES = {
    "uz": {
        "ijara": "🏠 Ijara shartnomasi",
        "oldi_sotdi": "🤝 Oldi-sotdi shartnomasi",
        "tilxat": "✍️ Tilxat",
        "ishonchnoma": "📋 Ishonchnoma",
        "nikoh": "💍 Nikoh shartnomasi",
        "qarz": "💰 Qarz shartnomasi",
    },
    "ru": {
        "ijara": "🏠 Договор аренды",
        "oldi_sotdi": "🤝 Договор купли-продажи",
        "tilxat": "✍️ Расписка",
        "ishonchnoma": "📋 Доверенность",
        "nikoh": "💍 Брачный договор",
        "qarz": "💰 Договор займа",
    },
    "en": {
        "ijara": "🏠 Lease agreement",
        "oldi_sotdi": "🤝 Sale agreement",
        "tilxat": "✍️ Receipt",
        "ishonchnoma": "📋 Power of attorney",
        "nikoh": "💍 Marriage contract",
        "qarz": "💰 Loan agreement",
    }
}

PRICE = 10000

def t(lang, key):
    return TEXTS.get(lang, TEXTS["uz"]).get(key, key)

def dn(lang, doc):
    return DOC_NAMES.get(lang, DOC_NAMES["uz"]).get(doc, doc)

# ─── KLAVIATURALAR ────────────────────────────────────────────────────────────

def main_menu(lang="uz"):
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=t(lang, "menu_docs"), callback_data="menu_docs")],
        [InlineKeyboardButton(text=t(lang, "menu_ai"), callback_data="menu_ai")],
        [InlineKeyboardButton(text=t(lang, "menu_info"), callback_data="menu_info")],
        [InlineKeyboardButton(text=t(lang, "menu_lang"), callback_data="menu_lang")],
    ])

def docs_menu(lang="uz"):
    rows = []
    for key in ["ijara", "oldi_sotdi", "tilxat", "ishonchnoma", "nikoh", "qarz"]:
        rows.append([InlineKeyboardButton(
            text=f"{dn(lang, key)} — 10,000",
            callback_data=f"doc_{key}"
        )])
    rows.append([InlineKeyboardButton(text=t(lang, "back"), callback_data="back_main")])
    return InlineKeyboardMarkup(inline_keyboard=rows)

def pay_menu(order_id, lang="uz"):
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=t(lang, "pay_btn"), callback_data=f"pay_{order_id}")],
        [InlineKeyboardButton(text=t(lang, "back"), callback_data="back_main")],
    ])

def confirm_menu(order_id, lang="uz"):
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=t(lang, "confirm_btn"), callback_data=f"confirm_{order_id}")],
        [InlineKeyboardButton(text=t(lang, "back"), callback_data="back_main")],
    ])

def cancel_menu(lang="uz"):
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=t(lang, "cancel"), callback_data="back_main")],
    ])

def lang_menu():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🇺🇿 O'zbek", callback_data="setlang_uz")],
        [InlineKeyboardButton(text="🇷🇺 Русский", callback_data="setlang_ru")],
        [InlineKeyboardButton(text="🇬🇧 English", callback_data="setlang_en")],
    ])

def admin_menu(lang="uz"):
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📊 Statistika", callback_data="admin_stats")],
        [InlineKeyboardButton(text="👥 Foydalanuvchilar", callback_data="admin_users")],
        [InlineKeyboardButton(text="📋 So'nggi buyurtmalar", callback_data="admin_orders")],
        [InlineKeyboardButton(text="📢 Xabar tarqatish", callback_data="admin_broadcast")],
        [InlineKeyboardButton(text=t(lang, "back"), callback_data="back_main")],
    ])

# ─── MA'LUMOTLAR BAZASI ───────────────────────────────────────────────────────

async def init_db():
    async with aiosqlite.connect(DB) as db:
        await db.execute("""CREATE TABLE IF NOT EXISTS users (
            user_id INTEGER PRIMARY KEY, username TEXT, full_name TEXT,
            language TEXT DEFAULT 'uz', joined_at TEXT DEFAULT CURRENT_TIMESTAMP)""")
        await db.execute("""CREATE TABLE IF NOT EXISTS orders (
            id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER,
            doc_type TEXT, amount INTEGER, status TEXT DEFAULT 'pending',
            data TEXT, file_path TEXT, created_at TEXT DEFAULT CURRENT_TIMESTAMP)""")
        await db.commit()

async def save_user(user_id, username, full_name):
    async with aiosqlite.connect(DB) as db:
        await db.execute("INSERT OR IGNORE INTO users (user_id, username, full_name) VALUES (?,?,?)",
                         (user_id, username, full_name))
        await db.commit()

async def get_lang(user_id):
    async with aiosqlite.connect(DB) as db:
        async with db.execute("SELECT language FROM users WHERE user_id=?", (user_id,)) as cur:
            row = await cur.fetchone()
            return row[0] if row else "uz"

async def set_lang(user_id, lang):
    async with aiosqlite.connect(DB) as db:
        await db.execute("UPDATE users SET language=? WHERE user_id=?", (lang, user_id))
        await db.commit()

async def create_order(user_id, doc_type, amount, data):
    async with aiosqlite.connect(DB) as db:
        cur = await db.execute(
            "INSERT INTO orders (user_id, doc_type, amount, data) VALUES (?,?,?,?)",
            (user_id, doc_type, amount, json.dumps(data, ensure_ascii=False)))
        await db.commit()
        return cur.lastrowid

async def get_order(order_id):
    async with aiosqlite.connect(DB) as db:
        async with db.execute("SELECT * FROM orders WHERE id=?", (order_id,)) as cur:
            row = await cur.fetchone()
            if row:
                return {"id": row[0], "user_id": row[1], "doc_type": row[2],
                        "amount": row[3], "status": row[4],
                        "data": json.loads(row[5] or "{}"), "file_path": row[6]}

async def update_order(order_id, status, file_path=None):
    async with aiosqlite.connect(DB) as db:
        await db.execute("UPDATE orders SET status=?, file_path=? WHERE id=?",
                         (status, file_path, order_id))
        await db.commit()

async def get_stats():
    async with aiosqlite.connect(DB) as db:
        async with db.execute("SELECT COUNT(*) FROM users") as c:
            users = (await c.fetchone())[0]
        async with db.execute("SELECT COUNT(*), SUM(amount) FROM orders WHERE status='paid'") as c:
            row = await c.fetchone()
            orders, revenue = row[0], row[1] or 0
        async with db.execute("SELECT COUNT(*) FROM users WHERE joined_at >= date('now','-1 day')") as c:
            new_today = (await c.fetchone())[0]
    return users, orders, revenue, new_today

async def get_recent_users(limit=10):
    async with aiosqlite.connect(DB) as db:
        async with db.execute("SELECT user_id, username, full_name, joined_at FROM users ORDER BY joined_at DESC LIMIT ?", (limit,)) as cur:
            return await cur.fetchall()

async def get_recent_orders(limit=10):
    async with aiosqlite.connect(DB) as db:
        async with db.execute("SELECT id, user_id, doc_type, amount, status, created_at FROM orders ORDER BY created_at DESC LIMIT ?", (limit,)) as cur:
            return await cur.fetchall()

async def get_all_users():
    async with aiosqlite.connect(DB) as db:
        async with db.execute("SELECT user_id FROM users") as cur:
            return [r[0] for r in await cur.fetchall()]

# ─── HUJJAT GENERATSIYA ───────────────────────────────────────────────────────

def make_doc(doc_type, data, order_id):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    sana = datetime.now().strftime('%d.%m.%Y')

    def heading(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text).bold = True
        return p

    def line(text):
        doc.add_paragraph(text)

    if doc_type == "ijara":
        heading("IJARA SHARTNOMASI")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','Toshkent')}")
        line(f"Ijara beruvchi: {data.get('beruvchi_fio','___')} (pasport: {data.get('beruvchi_pasport','___')})")
        line(f"Ijara oluvchi: {data.get('oluvchi_fio','___')} (pasport: {data.get('oluvchi_pasport','___')})")
        line(f"Manzil: {data.get('manzil','___')}")
        line(f"Oylik ijara haqi: {data.get('narx','___')} so'm")
        line(f"Muddat: {data.get('muddat','___')} oy")
        line("\nImzo: __________          Imzo: __________")

    elif doc_type == "oldi_sotdi":
        heading("OLDI-SOTDI SHARTNOMASI")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','Toshkent')}")
        line(f"Sotuvchi: {data.get('sotuvchi_fio','___')} (pasport: {data.get('sotuvchi_pasport','___')})")
        line(f"Xaridor: {data.get('xaridor_fio','___')} (pasport: {data.get('xaridor_pasport','___')})")
        line(f"Mulk: {data.get('tovar','___')}")
        line(f"Narxi: {data.get('narx','___')} so'm")
        line("\nImzo: __________          Imzo: __________")

    elif doc_type == "tilxat":
        heading("TILXAT")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','Toshkent')}")
        line(f"Men, {data.get('fio','___')}, pasport: {data.get('pasport','___')}, {data.get('manzil','___')} manzilida yashayman.")
        line(f"Tasdiqlayman: {data.get('mazmun','___')}")
        line(f"Miqdori: {data.get('miqdor','___')} so'm")
        line("\nImzo: __________")

    elif doc_type == "ishonchnoma":
        heading("ISHONCHNOMA")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','Toshkent')}")
        line(f"Men, {data.get('beruvchi_fio','___')} (pasport: {data.get('beruvchi_pasport','___')}),")
        line(f"{data.get('oluvchi_fio','___')} (pasport: {data.get('oluvchi_pasport','___')}) ga quyidagi vakolatlarni beraman:")
        line(f"Vakolat: {data.get('vakolat','___')}")
        line(f"Muddat: {data.get('muddat','___')}")
        line("\nImzo: __________")

    elif doc_type == "nikoh":
        heading("NIKOH SHARTNOMASI")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','Toshkent')}")
        line(f"Er: {data.get('er_fio','___')} (pasport: {data.get('er_pasport','___')})")
        line(f"Xotin: {data.get('xotin_fio','___')} (pasport: {data.get('xotin_pasport','___')})")
        line(f"Mulkiy shartlar: {data.get('shartlar','___')}")
        line("\nImzo: __________          Imzo: __________")

    elif doc_type == "qarz":
        heading("QARZ SHARTNOMASI")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','Toshkent')}")
        line(f"Qarz beruvchi: {data.get('beruvchi_fio','___')} (pasport: {data.get('beruvchi_pasport','___')})")
        line(f"Qarz oluvchi: {data.get('oluvchi_fio','___')} (pasport: {data.get('oluvchi_pasport','___')})")
        line(f"Qarz miqdori: {data.get('miqdor','___')} so'm")
        line(f"Qaytarish muddati: {data.get('muddat','___')}")
        line(f"Foiz: {data.get('foiz','0')}%")
        line("\nImzo: __________          Imzo: __________")

    path = f"{DOCS_DIR}/{doc_type}_{order_id}.docx"
    doc.save(path)
    return path

# ─── FSM STATELARI ───────────────────────────────────────────────────────────

class Ijara(StatesGroup):
    beruvchi_fio = State(); beruvchi_pasport = State()
    oluvchi_fio = State(); oluvchi_pasport = State()
    manzil = State(); narx = State(); muddat = State(); shahar = State()

class OldiSotdi(StatesGroup):
    sotuvchi_fio = State(); sotuvchi_pasport = State()
    xaridor_fio = State(); xaridor_pasport = State()
    tovar = State(); narx = State(); shahar = State()

class Tilxat(StatesGroup):
    fio = State(); pasport = State(); manzil = State()
    mazmun = State(); miqdor = State(); shahar = State()

class Ishonchnoma(StatesGroup):
    beruvchi_fio = State(); beruvchi_pasport = State()
    oluvchi_fio = State(); oluvchi_pasport = State()
    vakolat = State(); muddat = State(); shahar = State()

class Nikoh(StatesGroup):
    er_fio = State(); er_pasport = State()
    xotin_fio = State(); xotin_pasport = State()
    shartlar = State(); shahar = State()

class Qarz(StatesGroup):
    beruvchi_fio = State(); beruvchi_pasport = State()
    oluvchi_fio = State(); oluvchi_pasport = State()
    miqdor = State(); muddat = State(); foiz = State(); shahar = State()

class AI(StatesGroup):
    savol = State()

class Broadcast(StatesGroup):
    xabar = State()

# ─── BOT ─────────────────────────────────────────────────────────────────────

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher(storage=MemoryStorage())

async def get_user_lang(user_id):
    return await get_lang(user_id)

@dp.message(CommandStart())
async def start(message: Message):
    await save_user(message.from_user.id, message.from_user.username or "", message.from_user.full_name or "")
    lang = await get_user_lang(message.from_user.id)
    await message.answer(t(lang, "welcome"), parse_mode="HTML", reply_markup=main_menu(lang))

@dp.callback_query(F.data == "back_main")
async def back_main(cb: CallbackQuery, state: FSMContext):
    await state.clear()
    lang = await get_user_lang(cb.from_user.id)
    await cb.message.edit_text(t(lang, "welcome"), parse_mode="HTML", reply_markup=main_menu(lang))

@dp.callback_query(F.data == "menu_info")
async def info(cb: CallbackQuery):
    lang = await get_user_lang(cb.from_user.id)
    await cb.message.edit_text(t(lang, "info"), parse_mode="HTML", reply_markup=main_menu(lang))

@dp.callback_query(F.data == "menu_docs")
async def docs(cb: CallbackQuery):
    lang = await get_user_lang(cb.from_user.id)
    await cb.message.edit_text(t(lang, "choose_doc"), parse_mode="HTML", reply_markup=docs_menu(lang))

@dp.callback_query(F.data == "menu_lang")
async def choose_lang(cb: CallbackQuery):
    lang = await get_user_lang(cb.from_user.id)
    await cb.message.edit_text(t(lang, "choose_lang"), reply_markup=lang_menu())

@dp.callback_query(F.data.startswith("setlang_"))
async def set_language(cb: CallbackQuery):
    new_lang = cb.data.split("_")[1]
    await set_lang(cb.from_user.id, new_lang)
    await cb.message.edit_text(t(new_lang, "lang_saved"), reply_markup=main_menu(new_lang))

# ─── IJARA ───────────────────────────────────────────────────────────────────

@dp.callback_query(F.data == "doc_ijara")
async def ijara_start(cb: CallbackQuery, state: FSMContext):
    lang = await get_user_lang(cb.from_user.id)
    await state.set_state(Ijara.beruvchi_fio)
    await cb.message.edit_text(f"🏠 <b>{dn(lang,'ijara')}</b>\n\n1️⃣ Ijara beruvchining to'liq ismi:", parse_mode="HTML", reply_markup=cancel_menu(lang))

@dp.message(Ijara.beruvchi_fio)
async def i1(m: Message, state: FSMContext):
    await state.update_data(beruvchi_fio=m.text); await state.set_state(Ijara.beruvchi_pasport)
    await m.answer("2️⃣ Ijara beruvchining pasporti:")

@dp.message(Ijara.beruvchi_pasport)
async def i2(m: Message, state: FSMContext):
    await state.update_data(beruvchi_pasport=m.text); await state.set_state(Ijara.oluvchi_fio)
    await m.answer("3️⃣ Ijara oluvchining to'liq ismi:")

@dp.message(Ijara.oluvchi_fio)
async def i3(m: Message, state: FSMContext):
    await state.update_data(oluvchi_fio=m.text); await state.set_state(Ijara.oluvchi_pasport)
    await m.answer("4️⃣ Ijara oluvchining pasporti:")

@dp.message(Ijara.oluvchi_pasport)
async def i4(m: Message, state: FSMContext):
    await state.update_data(oluvchi_pasport=m.text); await state.set_state(Ijara.manzil)
    await m.answer("5️⃣ Mulkning to'liq manzili:")

@dp.message(Ijara.manzil)
async def i5(m: Message, state: FSMContext):
    await state.update_data(manzil=m.text); await state.set_state(Ijara.narx)
    await m.answer("6️⃣ Oylik ijara haqi (so'mda):")

@dp.message(Ijara.narx)
async def i6(m: Message, state: FSMContext):
    await state.update_data(narx=m.text); await state.set_state(Ijara.muddat)
    await m.answer("7️⃣ Ijara muddati (oyda):")

@dp.message(Ijara.muddat)
async def i7(m: Message, state: FSMContext):
    await state.update_data(muddat=m.text); await state.set_state(Ijara.shahar)
    await m.answer("8️⃣ Shahar:")

@dp.message(Ijara.shahar)
async def i8(m: Message, state: FSMContext):
    await state.update_data(shahar=m.text)
    data = await state.get_data(); await state.clear()
    lang = await get_user_lang(m.from_user.id)
    order_id = await create_order(m.from_user.id, "ijara", PRICE, data)
    await m.answer(f"✅ <b>Ma'lumotlar qabul qilindi!</b>\n\n• Beruvchi: {data['beruvchi_fio']}\n• Oluvchi: {data['oluvchi_fio']}\n• Manzil: {data['manzil']}\n• Narx: {data['narx']} so'm\n\n💳 To'lov qiling:", parse_mode="HTML", reply_markup=pay_menu(order_id, lang))

# ─── OLDI-SOTDI ───────────────────────────────────────────────────────────────

@dp.callback_query(F.data == "doc_oldi_sotdi")
async def os_start(cb: CallbackQuery, state: FSMContext):
    lang = await get_user_lang(cb.from_user.id)
    await state.set_state(OldiSotdi.sotuvchi_fio)
    await cb.message.edit_text(f"🤝 <b>{dn(lang,'oldi_sotdi')}</b>\n\n1️⃣ Sotuvchining to'liq ismi:", parse_mode="HTML", reply_markup=cancel_menu(lang))

@dp.message(OldiSotdi.sotuvchi_fio)
async def os1(m: Message, state: FSMContext):
    await state.update_data(sotuvchi_fio=m.text); await state.set_state(OldiSotdi.sotuvchi_pasport)
    await m.answer("2️⃣ Sotuvchining pasporti:")

@dp.message(OldiSotdi.sotuvchi_pasport)
async def os2(m: Message, state: FSMContext):
    await state.update_data(sotuvchi_pasport=m.text); await state.set_state(OldiSotdi.xaridor_fio)
    await m.answer("3️⃣ Xaridorning to'liq ismi:")

@dp.message(OldiSotdi.xaridor_fio)
async def os3(m: Message, state: FSMContext):
    await state.update_data(xaridor_fio=m.text); await state.set_state(OldiSotdi.xaridor_pasport)
    await m.answer("4️⃣ Xaridorning pasporti:")

@dp.message(OldiSotdi.xaridor_pasport)
async def os4(m: Message, state: FSMContext):
    await state.update_data(xaridor_pasport=m.text); await state.set_state(OldiSotdi.tovar)
    await m.answer("5️⃣ Sotilayotgan mulk tavsifi:")

@dp.message(OldiSotdi.tovar)
async def os5(m: Message, state: FSMContext):
    await state.update_data(tovar=m.text); await state.set_state(OldiSotdi.narx)
    await m.answer("6️⃣ Mulk narxi (so'mda):")

@dp.message(OldiSotdi.narx)
async def os6(m: Message, state: FSMContext):
    await state.update_data(narx=m.text); await state.set_state(OldiSotdi.shahar)
    await m.answer("7️⃣ Shahar:")

@dp.message(OldiSotdi.shahar)
async def os7(m: Message, state: FSMContext):
    await state.update_data(shahar=m.text)
    data = await state.get_data(); await state.clear()
    lang = await get_user_lang(m.from_user.id)
    order_id = await create_order(m.from_user.id, "oldi_sotdi", PRICE, data)
    await m.answer(f"✅ <b>Ma'lumotlar qabul qilindi!</b>\n\n• Sotuvchi: {data['sotuvchi_fio']}\n• Xaridor: {data['xaridor_fio']}\n• Mulk: {data['tovar']}\n• Narx: {data['narx']} so'm\n\n💳 To'lov qiling:", parse_mode="HTML", reply_markup=pay_menu(order_id, lang))

# ─── TILXAT ───────────────────────────────────────────────────────────────────

@dp.callback_query(F.data == "doc_tilxat")
async def tx_start(cb: CallbackQuery, state: FSMContext):
    lang = await get_user_lang(cb.from_user.id)
    await state.set_state(Tilxat.fio)
    await cb.message.edit_text(f"✍️ <b>{dn(lang,'tilxat')}</b>\n\n1️⃣ To'liq ismi familiyangiz:", parse_mode="HTML", reply_markup=cancel_menu(lang))

@dp.message(Tilxat.fio)
async def tx1(m: Message, state: FSMContext):
    await state.update_data(fio=m.text); await state.set_state(Tilxat.pasport)
    await m.answer("2️⃣ Pasport seriyasi va raqami:")

@dp.message(Tilxat.pasport)
async def tx2(m: Message, state: FSMContext):
    await state.update_data(pasport=m.text); await state.set_state(Tilxat.manzil)
    await m.answer("3️⃣ Yashash manzili:")

@dp.message(Tilxat.manzil)
async def tx3(m: Message, state: FSMContext):
    await state.update_data(manzil=m.text); await state.set_state(Tilxat.mazmun)
    await m.answer("4️⃣ Tilxat mazmuni:")

@dp.message(Tilxat.mazmun)
async def tx4(m: Message, state: FSMContext):
    await state.update_data(mazmun=m.text); await state.set_state(Tilxat.miqdor)
    await m.answer("5️⃣ Miqdori (so'mda):")

@dp.message(Tilxat.miqdor)
async def tx5(m: Message, state: FSMContext):
    await state.update_data(miqdor=m.text); await state.set_state(Tilxat.shahar)
    await m.answer("6️⃣ Shahar:")

@dp.message(Tilxat.shahar)
async def tx6(m: Message, state: FSMContext):
    await state.update_data(shahar=m.text)
    data = await state.get_data(); await state.clear()
    lang = await get_user_lang(m.from_user.id)
    order_id = await create_order(m.from_user.id, "tilxat", PRICE, data)
    await m.answer(f"✅ <b>Ma'lumotlar qabul qilindi!</b>\n\n• Muallif: {data['fio']}\n• Miqdor: {data['miqdor']} so'm\n\n💳 To'lov qiling:", parse_mode="HTML", reply_markup=pay_menu(order_id, lang))

# ─── ISHONCHNOMA ──────────────────────────────────────────────────────────────

@dp.callback_query(F.data == "doc_ishonchnoma")
async def ish_start(cb: CallbackQuery, state: FSMContext):
    lang = await get_user_lang(cb.from_user.id)
    await state.set_state(Ishonchnoma.beruvchi_fio)
    await cb.message.edit_text(f"📋 <b>{dn(lang,'ishonchnoma')}</b>\n\n1️⃣ Ishonchnoma beruvchining to'liq ismi:", parse_mode="HTML", reply_markup=cancel_menu(lang))

@dp.message(Ishonchnoma.beruvchi_fio)
async def ish1(m: Message, state: FSMContext):
    await state.update_data(beruvchi_fio=m.text); await state.set_state(Ishonchnoma.beruvchi_pasport)
    await m.answer("2️⃣ Beruvchining pasporti:")

@dp.message(Ishonchnoma.beruvchi_pasport)
async def ish2(m: Message, state: FSMContext):
    await state.update_data(beruvchi_pasport=m.text); await state.set_state(Ishonchnoma.oluvchi_fio)
    await m.answer("3️⃣ Ishonchnoma oluvchining to'liq ismi:")

@dp.message(Ishonchnoma.oluvchi_fio)
async def ish3(m: Message, state: FSMContext):
    await state.update_data(oluvchi_fio=m.text); await state.set_state(Ishonchnoma.oluvchi_pasport)
    await m.answer("4️⃣ Oluvchining pasporti:")

@dp.message(Ishonchnoma.oluvchi_pasport)
async def ish4(m: Message, state: FSMContext):
    await state.update_data(oluvchi_pasport=m.text); await state.set_state(Ishonchnoma.vakolat)
    await m.answer("5️⃣ Beriladigan vakolat (nima qilish huquqi beriladi):")

@dp.message(Ishonchnoma.vakolat)
async def ish5(m: Message, state: FSMContext):
    await state.update_data(vakolat=m.text); await state.set_state(Ishonchnoma.muddat)
    await m.answer("6️⃣ Ishonchnoma muddati (masalan: 1 yil):")

@dp.message(Ishonchnoma.muddat)
async def ish6(m: Message, state: FSMContext):
    await state.update_data(muddat=m.text); await state.set_state(Ishonchnoma.shahar)
    await m.answer("7️⃣ Shahar:")

@dp.message(Ishonchnoma.shahar)
async def ish7(m: Message, state: FSMContext):
    await state.update_data(shahar=m.text)
    data = await state.get_data(); await state.clear()
    lang = await get_user_lang(m.from_user.id)
    order_id = await create_order(m.from_user.id, "ishonchnoma", PRICE, data)
    await m.answer(f"✅ <b>Ma'lumotlar qabul qilindi!</b>\n\n• Beruvchi: {data['beruvchi_fio']}\n• Oluvchi: {data['oluvchi_fio']}\n\n💳 To'lov qiling:", parse_mode="HTML", reply_markup=pay_menu(order_id, lang))

# ─── NIKOH SHARTNOMASI ────────────────────────────────────────────────────────

@dp.callback_query(F.data == "doc_nikoh")
async def nikoh_start(cb: CallbackQuery, state: FSMContext):
    lang = await get_user_lang(cb.from_user.id)
    await state.set_state(Nikoh.er_fio)
    await cb.message.edit_text(f"💍 <b>{dn(lang,'nikoh')}</b>\n\n1️⃣ Erning to'liq ismi:", parse_mode="HTML", reply_markup=cancel_menu(lang))

@dp.message(Nikoh.er_fio)
async def n1(m: Message, state: FSMContext):
    await state.update_data(er_fio=m.text); await state.set_state(Nikoh.er_pasport)
    await m.answer("2️⃣ Erning pasporti:")

@dp.message(Nikoh.er_pasport)
async def n2(m: Message, state: FSMContext):
    await state.update_data(er_pasport=m.text); await state.set_state(Nikoh.xotin_fio)
    await m.answer("3️⃣ Xotinning to'liq ismi:")

@dp.message(Nikoh.xotin_fio)
async def n3(m: Message, state: FSMContext):
    await state.update_data(xotin_fio=m.text); await state.set_state(Nikoh.xotin_pasport)
    await m.answer("4️⃣ Xotinning pasporti:")

@dp.message(Nikoh.xotin_pasport)
async def n4(m: Message, state: FSMContext):
    await state.update_data(xotin_pasport=m.text); await state.set_state(Nikoh.shartlar)
    await m.answer("5️⃣ Mulkiy shartlar (qanday kelishildi):")

@dp.message(Nikoh.shartlar)
async def n5(m: Message, state: FSMContext):
    await state.update_data(shartlar=m.text); await state.set_state(Nikoh.shahar)
    await m.answer("6️⃣ Shahar:")

@dp.message(Nikoh.shahar)
async def n6(m: Message, state: FSMContext):
    await state.update_data(shahar=m.text)
    data = await state.get_data(); await state.clear()
    lang = await get_user_lang(m.from_user.id)
    order_id = await create_order(m.from_user.id, "nikoh", PRICE, data)
    await m.answer(f"✅ <b>Ma'lumotlar qabul qilindi!</b>\n\n• Er: {data['er_fio']}\n• Xotin: {data['xotin_fio']}\n\n💳 To'lov qiling:", parse_mode="HTML", reply_markup=pay_menu(order_id, lang))

# ─── QARZ SHARTNOMASI ─────────────────────────────────────────────────────────

@dp.callback_query(F.data == "doc_qarz")
async def qarz_start(cb: CallbackQuery, state: FSMContext):
    lang = await get_user_lang(cb.from_user.id)
    await state.set_state(Qarz.beruvchi_fio)
    await cb.message.edit_text(f"💰 <b>{dn(lang,'qarz')}</b>\n\n1️⃣ Qarz beruvchining to'liq ismi:", parse_mode="HTML", reply_markup=cancel_menu(lang))

@dp.message(Qarz.beruvchi_fio)
async def q1(m: Message, state: FSMContext):
    await state.update_data(beruvchi_fio=m.text); await state.set_state(Qarz.beruvchi_pasport)
    await m.answer("2️⃣ Beruvchining pasporti:")

@dp.message(Qarz.beruvchi_pasport)
async def q2(m: Message, state: FSMContext):
    await state.update_data(beruvchi_pasport=m.text); await state.set_state(Qarz.oluvchi_fio)
    await m.answer("3️⃣ Qarz oluvchining to'liq ismi:")

@dp.message(Qarz.oluvchi_fio)
async def q3(m: Message, state: FSMContext):
    await state.update_data(oluvchi_fio=m.text); await state.set_state(Qarz.oluvchi_pasport)
    await m.answer("4️⃣ Oluvchining pasporti:")

@dp.message(Qarz.oluvchi_pasport)
async def q4(m: Message, state: FSMContext):
    await state.update_data(oluvchi_pasport=m.text); await state.set_state(Qarz.miqdor)
    await m.answer("5️⃣ Qarz miqdori (so'mda):")

@dp.message(Qarz.miqdor)
async def q5(m: Message, state: FSMContext):
    await state.update_data(miqdor=m.text); await state.set_state(Qarz.muddat)
    await m.answer("6️⃣ Qaytarish muddati (masalan: 3 oy):")

@dp.message(Qarz.muddat)
async def q6(m: Message, state: FSMContext):
    await state.update_data(muddat=m.text); await state.set_state(Qarz.foiz)
    await m.answer("7️⃣ Foiz stavkasi (0 bo'lsa 0 yozing):")

@dp.message(Qarz.foiz)
async def q7(m: Message, state: FSMContext):
    await state.update_data(foiz=m.text); await state.set_state(Qarz.shahar)
    await m.answer("8️⃣ Shahar:")

@dp.message(Qarz.shahar)
async def q8(m: Message, state: FSMContext):
    await state.update_data(shahar=m.text)
    data = await state.get_data(); await state.clear()
    lang = await get_user_lang(m.from_user.id)
    order_id = await create_order(m.from_user.id, "qarz", PRICE, data)
    await m.answer(f"✅ <b>Ma'lumotlar qabul qilindi!</b>\n\n• Beruvchi: {data['beruvchi_fio']}\n• Oluvchi: {data['oluvchi_fio']}\n• Miqdor: {data['miqdor']} so'm\n\n💳 To'lov qiling:", parse_mode="HTML", reply_markup=pay_menu(order_id, lang))

# ─── TO'LOV ───────────────────────────────────────────────────────────────────

@dp.callback_query(F.data.startswith("pay_"))
async def pay(cb: CallbackQuery):
    order_id = int(cb.data.split("_")[1])
    order = await get_order(order_id)
    lang = await get_user_lang(cb.from_user.id)
    if not order:
        await cb.answer("Buyurtma topilmadi!", show_alert=True); return
    await cb.message.edit_text(
        f"💳 <b>To'lov</b>\n\nBuyurtma #{order_id}\nMiqdor: {order['amount']:,} so'm\n\n<i>Demo rejim</i>",
        parse_mode="HTML", reply_markup=confirm_menu(order_id, lang))

@dp.callback_query(F.data.startswith("confirm_"))
async def confirm(cb: CallbackQuery):
    order_id = int(cb.data.split("_")[1])
    order = await get_order(order_id)
    lang = await get_user_lang(cb.from_user.id)
    if not order:
        await cb.answer("Buyurtma topilmadi!", show_alert=True); return
    await cb.message.edit_text(t(lang, "preparing"))
    file_path = make_doc(order["doc_type"], order["data"], order_id)
    await update_order(order_id, "paid", file_path)
    doc_file = FSInputFile(file_path)
    await cb.message.answer_document(doc_file, caption=t(lang, "paid_ok"), parse_mode="HTML", reply_markup=main_menu(lang))
    await cb.message.delete()

# ─── AI MASLAHAT ──────────────────────────────────────────────────────────────

@dp.callback_query(F.data == "menu_ai")
async def ai_start(cb: CallbackQuery, state: FSMContext):
    lang = await get_user_lang(cb.from_user.id)
    await state.set_state(AI.savol)
    await cb.message.edit_text(t(lang, "ai_start"), parse_mode="HTML", reply_markup=cancel_menu(lang))

@dp.message(AI.savol)
async def ai_answer(m: Message, state: FSMContext):
    await state.clear()
    lang = await get_user_lang(m.from_user.id)
    msg = await m.answer(t(lang, "thinking"))
    try:
        async with httpx.AsyncClient(timeout=30) as client:
            r = await client.post(
                "https://api.openai.com/v1/chat/completions",
                headers={"Authorization": f"Bearer {OPENAI_API_KEY}"},
                json={"model": "gpt-4", "messages": [
                    {"role": "system", "content": "Siz O'zbekiston yuridik maslahatchi botisiz. Faqat O'zbekiston qonunchiligiga asoslanib javob bering. Javob oxirida: '⚠️ Bu boshlang'ich maslahat, to'liq yordam uchun advokatga murojaat qiling' deb yozing."},
                    {"role": "user", "content": m.text}
                ], "max_tokens": 800}
            )
            answer = r.json()["choices"][0]["message"]["content"]
    except:
        answer = t(lang, "ai_error")
    await msg.delete()
    await m.answer(f"⚖️ <b>Huquqiy Maslahat:</b>\n\n{answer}", parse_mode="HTML", reply_markup=main_menu(lang))

# ─── ADMIN ────────────────────────────────────────────────────────────────────

@dp.message(Command("admin"))
async def admin_cmd(m: Message):
    lang = await get_user_lang(m.from_user.id)
    if m.from_user.id not in ADMIN_IDS:
        await m.answer(t(lang, "no_access")); return
    await m.answer(t(lang, "admin_title"), parse_mode="HTML", reply_markup=admin_menu(lang))

@dp.callback_query(F.data == "admin_stats")
async def stats(cb: CallbackQuery):
    if cb.from_user.id not in ADMIN_IDS: return
    users, orders, revenue, new_today = await get_stats()
    await cb.message.edit_text(
        f"📊 <b>Statistika</b>\n\n"
        f"👥 Jami foydalanuvchilar: <b>{users}</b>\n"
        f"🆕 Bugun qo'shildi: <b>{new_today}</b>\n"
        f"📄 Jami buyurtmalar: <b>{orders}</b>\n"
        f"💰 Jami daromad: <b>{revenue:,} so'm</b>",
        parse_mode="HTML", reply_markup=admin_menu())

@dp.callback_query(F.data == "admin_users")
async def admin_users(cb: CallbackQuery):
    if cb.from_user.id not in ADMIN_IDS: return
    users = await get_recent_users(10)
    text = "👥 <b>So'nggi 10 foydalanuvchi:</b>\n\n"
    for u in users:
        text += f"• {u[2]} (@{u[1] or 'nomsiz'}) — {u[3][:10]}\n"
    await cb.message.edit_text(text, parse_mode="HTML", reply_markup=admin_menu())

@dp.callback_query(F.data == "admin_orders")
async def admin_orders(cb: CallbackQuery):
    if cb.from_user.id not in ADMIN_IDS: return
    orders = await get_recent_orders(10)
    text = "📋 <b>So'nggi 10 buyurtma:</b>\n\n"
    for o in orders:
        status_icon = "✅" if o[4] == "paid" else "⏳"
        text += f"{status_icon} #{o[0]} | {o[2]} | {o[3]:,} so'm | {o[5][:10]}\n"
    await cb.message.edit_text(text, parse_mode="HTML", reply_markup=admin_menu())

@dp.callback_query(F.data == "admin_broadcast")
async def broadcast_start(cb: CallbackQuery, state: FSMContext):
    if cb.from_user.id not in ADMIN_IDS: return
    lang = await get_user_lang(cb.from_user.id)
    await state.set_state(Broadcast.xabar)
    await cb.message.edit_text(t(lang, "broadcast_ask"))

@dp.message(Broadcast.xabar)
async def broadcast_send(m: Message, state: FSMContext):
    if m.from_user.id not in ADMIN_IDS: return
    await state.clear()
    lang = await get_user_lang(m.from_user.id)
    users = await get_all_users()
    sent = 0
    for uid in users:
        try:
            await bot.send_message(uid, f"📢 {m.text}")
            sent += 1
        except:
            pass
    await m.answer(t(lang, "broadcast_done").format(n=sent), reply_markup=admin_menu(lang))

# ─── ISHGA TUSHIRISH ──────────────────────────────────────────────────────────

async def main():
    await init_db()
    logging.info("Bot v3 ishga tushdi!")
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())

# ═══════════════════════════════════════════════════════════════════════════════
# YANGI HUJJATLAR — 12 TA QO'SHIMCHA
# ═══════════════════════════════════════════════════════════════════════════════

# ─── MEHNAT SHARTNOMASI ───────────────────────────────────────────────────────

class Mehnat(StatesGroup):
    ish_beruvchi = State(); xodim_fio = State(); xodim_pasport = State()
    lavozim = State(); oylik = State(); muddat = State(); shahar = State()

@dp.callback_query(F.data == "doc_mehnat")
async def mehnat_start(cb: CallbackQuery, state: FSMContext):
    lang = await get_user_lang(cb.from_user.id)
    await state.set_state(Mehnat.ish_beruvchi)
    await cb.message.edit_text("📝 <b>Mehnat Shartnomasi</b>\n\n1️⃣ Ish beruvchi tashkilot/shaxs nomi:", parse_mode="HTML", reply_markup=cancel_menu(lang))

@dp.message(Mehnat.ish_beruvchi)
async def meh1(m: Message, state: FSMContext):
    await state.update_data(ish_beruvchi=m.text); await state.set_state(Mehnat.xodim_fio)
    await m.answer("2️⃣ Xodimning to'liq ismi:")

@dp.message(Mehnat.xodim_fio)
async def meh2(m: Message, state: FSMContext):
    await state.update_data(xodim_fio=m.text); await state.set_state(Mehnat.xodim_pasport)
    await m.answer("3️⃣ Xodimning pasporti:")

@dp.message(Mehnat.xodim_pasport)
async def meh3(m: Message, state: FSMContext):
    await state.update_data(xodim_pasport=m.text); await state.set_state(Mehnat.lavozim)
    await m.answer("4️⃣ Lavozimi (masalan: Dasturchi):")

@dp.message(Mehnat.lavozim)
async def meh4(m: Message, state: FSMContext):
    await state.update_data(lavozim=m.text); await state.set_state(Mehnat.oylik)
    await m.answer("5️⃣ Oylik maosh (so'mda):")

@dp.message(Mehnat.oylik)
async def meh5(m: Message, state: FSMContext):
    await state.update_data(oylik=m.text); await state.set_state(Mehnat.muddat)
    await m.answer("6️⃣ Shartnoma muddati (masalan: 1 yil yoki Muddatsiz):")

@dp.message(Mehnat.muddat)
async def meh6(m: Message, state: FSMContext):
    await state.update_data(muddat=m.text); await state.set_state(Mehnat.shahar)
    await m.answer("7️⃣ Shahar:")

@dp.message(Mehnat.shahar)
async def meh7(m: Message, state: FSMContext):
    await state.update_data(shahar=m.text)
    data = await state.get_data(); await state.clear()
    lang = await get_user_lang(m.from_user.id)
    order_id = await create_order(m.from_user.id, "mehnat", PRICE, data)
    await m.answer(f"✅ <b>Ma'lumotlar qabul qilindi!</b>\n\n• Xodim: {data['xodim_fio']}\n• Lavozim: {data['lavozim']}\n• Oylik: {data['oylik']} so'm\n\n💳 To'lov qiling:", parse_mode="HTML", reply_markup=pay_menu(order_id, lang))

# ─── PUDRAT SHARTNOMASI ───────────────────────────────────────────────────────

class Pudrat(StatesGroup):
    buyurtmachi = State(); pudratchi = State(); pudratchi_pasport = State()
    ish_tavsifi = State(); narx = State(); muddat = State(); shahar = State()

@dp.callback_query(F.data == "doc_pudrat")
async def pudrat_start(cb: CallbackQuery, state: FSMContext):
    lang = await get_user_lang(cb.from_user.id)
    await state.set_state(Pudrat.buyurtmachi)
    await cb.message.edit_text("🏗️ <b>Pudrat Shartnomasi</b>\n\n1️⃣ Buyurtmachi (ism yoki tashkilot):", parse_mode="HTML", reply_markup=cancel_menu(lang))

@dp.message(Pudrat.buyurtmachi)
async def pud1(m: Message, state: FSMContext):
    await state.update_data(buyurtmachi=m.text); await state.set_state(Pudrat.pudratchi)
    await m.answer("2️⃣ Pudratchi (ism yoki tashkilot):")

@dp.message(Pudrat.pudratchi)
async def pud2(m: Message, state: FSMContext):
    await state.update_data(pudratchi=m.text); await state.set_state(Pudrat.pudratchi_pasport)
    await m.answer("3️⃣ Pudratchi pasporti yoki STIR:")

@dp.message(Pudrat.pudratchi_pasport)
async def pud3(m: Message, state: FSMContext):
    await state.update_data(pudratchi_pasport=m.text); await state.set_state(Pudrat.ish_tavsifi)
    await m.answer("4️⃣ Bajariladigan ish tavsifi:")

@dp.message(Pudrat.ish_tavsifi)
async def pud4(m: Message, state: FSMContext):
    await state.update_data(ish_tavsifi=m.text); await state.set_state(Pudrat.narx)
    await m.answer("5️⃣ Ish narxi (so'mda):")

@dp.message(Pudrat.narx)
async def pud5(m: Message, state: FSMContext):
    await state.update_data(narx=m.text); await state.set_state(Pudrat.muddat)
    await m.answer("6️⃣ Bajarish muddati:")

@dp.message(Pudrat.muddat)
async def pud6(m: Message, state: FSMContext):
    await state.update_data(muddat=m.text); await state.set_state(Pudrat.shahar)
    await m.answer("7️⃣ Shahar:")

@dp.message(Pudrat.shahar)
async def pud7(m: Message, state: FSMContext):
    await state.update_data(shahar=m.text)
    data = await state.get_data(); await state.clear()
    lang = await get_user_lang(m.from_user.id)
    order_id = await create_order(m.from_user.id, "pudrat", PRICE, data)
    await m.answer(f"✅ <b>Ma'lumotlar qabul qilindi!</b>\n\n• Buyurtmachi: {data['buyurtmachi']}\n• Pudratchi: {data['pudratchi']}\n• Narx: {data['narx']} so'm\n\n💳 To'lov qiling:", parse_mode="HTML", reply_markup=pay_menu(order_id, lang))

# ─── ALIMENT SHARTNOMASI ──────────────────────────────────────────────────────

class Aliment(StatesGroup):
    tolovchi = State(); tolovchi_pasport = State()
    oluvchi = State(); bola_ismi = State()
    miqdor = State(); muddat = State(); shahar = State()

@dp.callback_query(F.data == "doc_aliment")
async def aliment_start(cb: CallbackQuery, state: FSMContext):
    lang = await get_user_lang(cb.from_user.id)
    await state.set_state(Aliment.tolovchi)
    await cb.message.edit_text("👨‍👩‍👧 <b>Aliment Shartnomasi</b>\n\n1️⃣ Aliment to'lovchining to'liq ismi:", parse_mode="HTML", reply_markup=cancel_menu(lang))

@dp.message(Aliment.tolovchi)
async def al1(m: Message, state: FSMContext):
    await state.update_data(tolovchi=m.text); await state.set_state(Aliment.tolovchi_pasport)
    await m.answer("2️⃣ To'lovchining pasporti:")

@dp.message(Aliment.tolovchi_pasport)
async def al2(m: Message, state: FSMContext):
    await state.update_data(tolovchi_pasport=m.text); await state.set_state(Aliment.oluvchi)
    await m.answer("3️⃣ Aliment oluvchining to'liq ismi:")

@dp.message(Aliment.oluvchi)
async def al3(m: Message, state: FSMContext):
    await state.update_data(oluvchi=m.text); await state.set_state(Aliment.bola_ismi)
    await m.answer("4️⃣ Farzand(lar) ismi va tug'ilgan yili:")

@dp.message(Aliment.bola_ismi)
async def al4(m: Message, state: FSMContext):
    await state.update_data(bola_ismi=m.text); await state.set_state(Aliment.miqdor)
    await m.answer("5️⃣ Oylik aliment miqdori (so'mda):")

@dp.message(Aliment.miqdor)
async def al5(m: Message, state: FSMContext):
    await state.update_data(miqdor=m.text); await state.set_state(Aliment.muddat)
    await m.answer("6️⃣ To'lov muddati (masalan: Bola 18 yoshga to'lguncha):")

@dp.message(Aliment.muddat)
async def al6(m: Message, state: FSMContext):
    await state.update_data(muddat=m.text); await state.set_state(Aliment.shahar)
    await m.answer("7️⃣ Shahar:")

@dp.message(Aliment.shahar)
async def al7(m: Message, state: FSMContext):
    await state.update_data(shahar=m.text)
    data = await state.get_data(); await state.clear()
    lang = await get_user_lang(m.from_user.id)
    order_id = await create_order(m.from_user.id, "aliment", PRICE, data)
    await m.answer(f"✅ <b>Ma'lumotlar qabul qilindi!</b>\n\n• To'lovchi: {data['tolovchi']}\n• Farzand: {data['bola_ismi']}\n• Miqdor: {data['miqdor']} so'm\n\n💳 To'lov qiling:", parse_mode="HTML", reply_markup=pay_menu(order_id, lang))

# ─── HAMKORLIK SHARTNOMASI ────────────────────────────────────────────────────

class Hamkorlik(StatesGroup):
    tomon1 = State(); tomon1_pasport = State()
    tomon2 = State(); tomon2_pasport = State()
    loyiha = State(); foyda_taqsim = State(); muddat = State(); shahar = State()

@dp.callback_query(F.data == "doc_hamkorlik")
async def hamkorlik_start(cb: CallbackQuery, state: FSMContext):
    lang = await get_user_lang(cb.from_user.id)
    await state.set_state(Hamkorlik.tomon1)
    await cb.message.edit_text("🏢 <b>Hamkorlik Shartnomasi</b>\n\n1️⃣ 1-tomon to'liq ismi yoki tashkilot:", parse_mode="HTML", reply_markup=cancel_menu(lang))

@dp.message(Hamkorlik.tomon1)
async def ham1(m: Message, state: FSMContext):
    await state.update_data(tomon1=m.text); await state.set_state(Hamkorlik.tomon1_pasport)
    await m.answer("2️⃣ 1-tomon pasporti yoki STIR:")

@dp.message(Hamkorlik.tomon1_pasport)
async def ham2(m: Message, state: FSMContext):
    await state.update_data(tomon1_pasport=m.text); await state.set_state(Hamkorlik.tomon2)
    await m.answer("3️⃣ 2-tomon to'liq ismi yoki tashkilot:")

@dp.message(Hamkorlik.tomon2)
async def ham3(m: Message, state: FSMContext):
    await state.update_data(tomon2=m.text); await state.set_state(Hamkorlik.tomon2_pasport)
    await m.answer("4️⃣ 2-tomon pasporti yoki STIR:")

@dp.message(Hamkorlik.tomon2_pasport)
async def ham4(m: Message, state: FSMContext):
    await state.update_data(tomon2_pasport=m.text); await state.set_state(Hamkorlik.loyiha)
    await m.answer("5️⃣ Hamkorlik loyihasi/maqsadi:")

@dp.message(Hamkorlik.loyiha)
async def ham5(m: Message, state: FSMContext):
    await state.update_data(loyiha=m.text); await state.set_state(Hamkorlik.foyda_taqsim)
    await m.answer("6️⃣ Foyda taqsimoti (masalan: 50/50):")

@dp.message(Hamkorlik.foyda_taqsim)
async def ham6(m: Message, state: FSMContext):
    await state.update_data(foyda_taqsim=m.text); await state.set_state(Hamkorlik.muddat)
    await m.answer("7️⃣ Hamkorlik muddati:")

@dp.message(Hamkorlik.muddat)
async def ham7(m: Message, state: FSMContext):
    await state.update_data(muddat=m.text); await state.set_state(Hamkorlik.shahar)
    await m.answer("8️⃣ Shahar:")

@dp.message(Hamkorlik.shahar)
async def ham8(m: Message, state: FSMContext):
    await state.update_data(shahar=m.text)
    data = await state.get_data(); await state.clear()
    lang = await get_user_lang(m.from_user.id)
    order_id = await create_order(m.from_user.id, "hamkorlik", PRICE, data)
    await m.answer(f"✅ <b>Ma'lumotlar qabul qilindi!</b>\n\n• 1-tomon: {data['tomon1']}\n• 2-tomon: {data['tomon2']}\n• Loyiha: {data['loyiha']}\n\n💳 To'lov qiling:", parse_mode="HTML", reply_markup=pay_menu(order_id, lang))

# ─── DA'VO ARIZASI ────────────────────────────────────────────────────────────

class Davo(StatesGroup):
    ariza_beruvchi = State(); pasport = State()
    javobgar = State(); muammo = State(); talab = State(); shahar = State()

@dp.callback_query(F.data == "doc_davo")
async def davo_start(cb: CallbackQuery, state: FSMContext):
    lang = await get_user_lang(cb.from_user.id)
    await state.set_state(Davo.ariza_beruvchi)
    await cb.message.edit_text("⚖️ <b>Da'vo Arizasi</b>\n\n1️⃣ Ariza beruvchining to'liq ismi:", parse_mode="HTML", reply_markup=cancel_menu(lang))

@dp.message(Davo.ariza_beruvchi)
async def dav1(m: Message, state: FSMContext):
    await state.update_data(ariza_beruvchi=m.text); await state.set_state(Davo.pasport)
    await m.answer("2️⃣ Pasporti:")

@dp.message(Davo.pasport)
async def dav2(m: Message, state: FSMContext):
    await state.update_data(pasport=m.text); await state.set_state(Davo.javobgar)
    await m.answer("3️⃣ Javobgar (kim ustidan ariza):")

@dp.message(Davo.javobgar)
async def dav3(m: Message, state: FSMContext):
    await state.update_data(javobgar=m.text); await state.set_state(Davo.muammo)
    await m.answer("4️⃣ Muammo tavsifi (nima bo'ldi):")

@dp.message(Davo.muammo)
async def dav4(m: Message, state: FSMContext):
    await state.update_data(muammo=m.text); await state.set_state(Davo.talab)
    await m.answer("5️⃣ Talabingiz (nima so'rayapsiz):")

@dp.message(Davo.talab)
async def dav5(m: Message, state: FSMContext):
    await state.update_data(talab=m.text); await state.set_state(Davo.shahar)
    await m.answer("6️⃣ Shahar:")

@dp.message(Davo.shahar)
async def dav6(m: Message, state: FSMContext):
    await state.update_data(shahar=m.text)
    data = await state.get_data(); await state.clear()
    lang = await get_user_lang(m.from_user.id)
    order_id = await create_order(m.from_user.id, "davo", PRICE, data)
    await m.answer(f"✅ <b>Ma'lumotlar qabul qilindi!</b>\n\n• Ariza beruvchi: {data['ariza_beruvchi']}\n• Javobgar: {data['javobgar']}\n\n💳 To'lov qiling:", parse_mode="HTML", reply_markup=pay_menu(order_id, lang))

# ─── HUJJAT GENERATSIYA — YANGI TURLAR ───────────────────────────────────────

_original_make_doc = make_doc

def make_doc(doc_type, data, order_id):
    # Avvalgi turlar
    if doc_type in ["ijara", "oldi_sotdi", "tilxat", "ishonchnoma", "nikoh", "qarz"]:
        return _original_make_doc(doc_type, data, order_id)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    sana = datetime.now().strftime('%d.%m.%Y')

    def heading(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text).bold = True

    def line(text):
        doc.add_paragraph(text)

    if doc_type == "mehnat":
        heading("MEHNAT SHARTNOMASI")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','')}")
        line(f"Ish beruvchi: {data.get('ish_beruvchi','___')}")
        line(f"Xodim: {data.get('xodim_fio','___')} (pasport: {data.get('xodim_pasport','___')})")
        line(f"Lavozim: {data.get('lavozim','___')}")
        line(f"Oylik maosh: {data.get('oylik','___')} so'm")
        line(f"Shartnoma muddati: {data.get('muddat','___')}")
        line(f"\nIsh vaqti: Dushanba-Juma, 09:00-18:00")
        line(f"Ta'til: Yiliga 15 ish kuni")
        line("\nImzo: __________          Imzo: __________")

    elif doc_type == "pudrat":
        heading("PUDRAT SHARTNOMASI")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','')}")
        line(f"Buyurtmachi: {data.get('buyurtmachi','___')}")
        line(f"Pudratchi: {data.get('pudratchi','___')} ({data.get('pudratchi_pasport','___')})")
        line(f"Ish tavsifi: {data.get('ish_tavsifi','___')}")
        line(f"Ish narxi: {data.get('narx','___')} so'm")
        line(f"Bajarish muddati: {data.get('muddat','___')}")
        line("\nImzo: __________          Imzo: __________")

    elif doc_type == "aliment":
        heading("ALIMENT TO'LASH SHARTNOMASI")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','')}")
        line(f"To'lovchi: {data.get('tolovchi','___')} (pasport: {data.get('tolovchi_pasport','___')})")
        line(f"Oluvchi: {data.get('oluvchi','___')}")
        line(f"Farzand(lar): {data.get('bola_ismi','___')}")
        line(f"Oylik miqdor: {data.get('miqdor','___')} so'm")
        line(f"Muddat: {data.get('muddat','___')}")
        line("\nImzo: __________          Imzo: __________")

    elif doc_type == "hamkorlik":
        heading("HAMKORLIK SHARTNOMASI")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','')}")
        line(f"1-tomon: {data.get('tomon1','___')} ({data.get('tomon1_pasport','___')})")
        line(f"2-tomon: {data.get('tomon2','___')} ({data.get('tomon2_pasport','___')})")
        line(f"Loyiha: {data.get('loyiha','___')}")
        line(f"Foyda taqsimoti: {data.get('foyda_taqsim','___')}")
        line(f"Muddat: {data.get('muddat','___')}")
        line("\nImzo: __________          Imzo: __________")

    elif doc_type == "davo":
        heading("DA'VO ARIZASI")
        line(f"Sana: {sana}")
        line(f"Sudga: {data.get('shahar','___')} tumani sudiga")
        line(f"\nAriza beruvchi: {data.get('ariza_beruvchi','___')} (pasport: {data.get('pasport','___')})")
        line(f"Javobgar: {data.get('javobgar','___')}")
        line(f"\nMuammo: {data.get('muammo','___')}")
        line(f"\nTalabim: {data.get('talab','___')}")
        line(f"\nAriza beruvchi imzosi: __________")
        line(f"Sana: {sana}")

    path = f"{DOCS_DIR}/{doc_type}_{order_id}.docx"
    doc.save(path)
    return path

# ─── YANGI HUJJATLAR MENYUSI ─────────────────────────────────────────────────

@dp.callback_query(F.data == "menu_docs")
async def docs_new(cb: CallbackQuery):
    lang = await get_user_lang(cb.from_user.id)
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📁 Asosiy hujjatlar", callback_data="cat_asosiy")],
        [InlineKeyboardButton(text="📝 Ish hujjatlari", callback_data="cat_ish")],
        [InlineKeyboardButton(text="🏗️ Qurilish", callback_data="cat_qurilish")],
        [InlineKeyboardButton(text="👨‍👩‍👧 Oila", callback_data="cat_oila")],
        [InlineKeyboardButton(text="🏢 Biznes", callback_data="cat_biznes")],
        [InlineKeyboardButton(text="⚖️ Yuridik", callback_data="cat_yuridik")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="back_main")],
    ])
    await cb.message.edit_text("📂 <b>Kategoriyani tanlang:</b>", parse_mode="HTML", reply_markup=keyboard)

@dp.callback_query(F.data == "cat_asosiy")
async def cat_asosiy(cb: CallbackQuery):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🏠 Ijara shartnomasi", callback_data="doc_ijara")],
        [InlineKeyboardButton(text="🤝 Oldi-sotdi", callback_data="doc_oldi_sotdi")],
        [InlineKeyboardButton(text="✍️ Tilxat", callback_data="doc_tilxat")],
        [InlineKeyboardButton(text="📋 Ishonchnoma", callback_data="doc_ishonchnoma")],
        [InlineKeyboardButton(text="💍 Nikoh shartnomasi", callback_data="doc_nikoh")],
        [InlineKeyboardButton(text="💰 Qarz shartnomasi", callback_data="doc_qarz")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="menu_docs")],
    ])
    await cb.message.edit_text("📁 <b>Asosiy hujjatlar:</b>", parse_mode="HTML", reply_markup=keyboard)

@dp.callback_query(F.data == "cat_ish")
async def cat_ish(cb: CallbackQuery):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📝 Mehnat shartnomasi", callback_data="doc_mehnat")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="menu_docs")],
    ])
    await cb.message.edit_text("📝 <b>Ish hujjatlari:</b>", parse_mode="HTML", reply_markup=keyboard)

@dp.callback_query(F.data == "cat_qurilish")
async def cat_qurilish(cb: CallbackQuery):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🏗️ Pudrat shartnomasi", callback_data="doc_pudrat")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="menu_docs")],
    ])
    await cb.message.edit_text("🏗️ <b>Qurilish hujjatlari:</b>", parse_mode="HTML", reply_markup=keyboard)

@dp.callback_query(F.data == "cat_oila")
async def cat_oila(cb: CallbackQuery):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="👨‍👩‍👧 Aliment shartnomasi", callback_data="doc_aliment")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="menu_docs")],
    ])
    await cb.message.edit_text("👨‍👩‍👧 <b>Oila hujjatlari:</b>", parse_mode="HTML", reply_markup=keyboard)

@dp.callback_query(F.data == "cat_biznes")
async def cat_biznes(cb: CallbackQuery):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🏢 Hamkorlik shartnomasi", callback_data="doc_hamkorlik")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="menu_docs")],
    ])
    await cb.message.edit_text("🏢 <b>Biznes hujjatlari:</b>", parse_mode="HTML", reply_markup=keyboard)

@dp.callback_query(F.data == "cat_yuridik")
async def cat_yuridik(cb: CallbackQuery):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="⚖️ Da'vo arizasi", callback_data="doc_davo")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="menu_docs")],
    ])
    await cb.message.edit_text("⚖️ <b>Yuridik hujjatlar:</b>", parse_mode="HTML", reply_markup=keyboard)
